"""Generate WEBFLOW_PROGRAMMATIC_SEO_BRIEF.docx"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# Style defaults
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x0D, 0x0D, 0x0F)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# ── Title ──
p = doc.add_heading('Webflow Designer Brief', level=1)
p = doc.add_heading('Programmatic SEO Pages for kolo.xyz', level=2)

add_table(doc, ['Field', 'Value'], [
    ['For', 'Webflow Designer with Designer-level access to kolo.xyz'],
    ['From', 'SEO Agent App (kolo-seo.streamlit.app)'],
    ['Goal', 'Publish ~41 auto-generated landing pages targeting long-tail crypto card keywords'],
    ['Time', '~30 minutes of Designer work'],
    ['URL pattern', 'kolo.xyz/crypto-card/{slug} (e.g. kolo.xyz/crypto-card/uae)'],
])

# ── What You Receive ──
doc.add_heading('What You Will Receive From Us', level=2)
p = doc.add_paragraph('A ZIP folder called ')
p.add_run('pages/').bold = True
p.add_run(' containing:')

for item in [
    'cms_import.csv - Upload this to Webflow CMS (all pages in one file)',
    'manifest.json - Master list with all page data',
    'sitemap_fragment.xml - Sitemap entries for the new pages',
    'DESIGNER_BRIEF.md - This document',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('The CSV has all content pre-written \u2014 titles, descriptions, H1s, body content, FAQ, structured data. You don\u2019t need to write anything.')

# ── Step 1 ──
doc.add_heading('Step 1: Create CMS Collection (5 min)', level=2)
doc.add_paragraph('1. Open Webflow Designer \u2192 CMS \u2192 + New Collection')
doc.add_paragraph('2. Name: Crypto Card Pages')
doc.add_paragraph('3. In Collection Settings \u2192 URL Structure \u2192 set prefix to crypto-card')
p = doc.add_paragraph('   This makes pages appear at ')
p.add_run('kolo.xyz/crypto-card/{slug}').italic = True
doc.add_paragraph('4. Add these fields:')

add_table(doc, ['#', 'Field Name', 'Field Type', 'Required'], [
    ['1', 'Name', 'Plain Text', 'Yes (auto-created)'],
    ['2', 'Slug', 'Slug', 'Yes (auto-created)'],
    ['3', 'SEO Title', 'Plain Text', 'Yes'],
    ['4', 'SEO Description', 'Plain Text', 'Yes'],
    ['5', 'H1', 'Plain Text', 'Yes'],
    ['6', 'Hero Subtitle', 'Plain Text', 'Yes'],
    ['7', 'Body Content', 'Rich Text', 'Yes'],
    ['8', 'FAQ Content', 'Rich Text', ''],
    ['9', 'Language', 'Option (en, ru, it, es, pl)', 'Yes'],
    ['10', 'Country', 'Plain Text', ''],
    ['11', 'Primary Keyword', 'Plain Text', ''],
    ['12', 'JSON-LD Code', 'Plain Text', ''],
    ['13', 'Hreflang Tags', 'Plain Text', ''],
    ['14', 'Related Pages', 'Multi-reference \u2192 Crypto Card Pages', ''],
])

doc.add_paragraph('5. Save the collection')

# ── Step 2 ──
doc.add_heading('Step 2: Design the CMS Template Page (15 min)', level=2)
p = doc.add_paragraph('Go to ')
p.add_run('CMS Collection Pages').bold = True
p.add_run(' \u2192 click ')
p.add_run('Crypto Card Pages Template').bold = True

doc.add_paragraph('Build the layout with these sections top-to-bottom:')

doc.add_heading('Navigation', level=3)
doc.add_paragraph('Use the existing kolo.xyz nav component (copy from homepage)', style='List Bullet')

doc.add_heading('Breadcrumbs', level=3)
doc.add_paragraph('Simple text: Home > Crypto Card > {Country}', style='List Bullet')
doc.add_paragraph('Bind {Country} to the Country CMS field', style='List Bullet')

doc.add_heading('Hero Section', level=3)
doc.add_paragraph('Dark background (#0D0D0F)', style='List Bullet')
doc.add_paragraph('H1 element \u2192 bind to H1 CMS field', style='List Bullet')
doc.add_paragraph('Subtitle paragraph \u2192 bind to Hero Subtitle CMS field', style='List Bullet')
doc.add_paragraph('CTA button: "Get Your Kolo Card" \u2192 link to https://kolo.xyz', style='List Bullet')
doc.add_paragraph('Style: large white text, centered, padding 80px top/bottom', style='List Bullet')

doc.add_heading('Content Section', level=3)
doc.add_paragraph('Rich Text element \u2192 bind to Body Content CMS field', style='List Bullet')
doc.add_paragraph('Max width: 720px, centered', style='List Bullet')
doc.add_paragraph('Style: white text on dark background, comfortable line height (1.7)', style='List Bullet')

doc.add_heading('FAQ Section', level=3)
doc.add_paragraph('Rich Text element \u2192 bind to FAQ Content CMS field', style='List Bullet')
doc.add_paragraph('Style as accordion if possible, or simple Q&A format', style='List Bullet')

doc.add_heading('CTA Banner', level=3)
doc.add_paragraph('Full-width dark section with accent color (#6C5CE7)', style='List Bullet')
doc.add_paragraph('Text: "Ready to spend crypto anywhere?"', style='List Bullet')
doc.add_paragraph('Button: "Get Your Card" \u2192 https://kolo.xyz', style='List Bullet')

doc.add_heading('Related Pages', level=3)
doc.add_paragraph('Collection List \u2192 filter by same Language OR same Country', style='List Bullet')
doc.add_paragraph('Show 3-4 related pages as cards with title + country', style='List Bullet')
doc.add_paragraph('Or bind to Related Pages multi-reference field', style='List Bullet')

doc.add_heading('Footer', level=3)
doc.add_paragraph('Use the existing kolo.xyz footer component (copy from homepage)', style='List Bullet')

doc.add_heading('SEO Settings (on the template page)', level=3)
doc.add_paragraph('Title Tag \u2192 bind to SEO Title field', style='List Bullet')
doc.add_paragraph('Meta Description \u2192 bind to SEO Description field', style='List Bullet')
doc.add_paragraph('Open Graph Title \u2192 bind to SEO Title field', style='List Bullet')
doc.add_paragraph('Open Graph Description \u2192 bind to SEO Description field', style='List Bullet')

doc.add_heading('Custom Code (Head)', level=3)
doc.add_paragraph('Add two Embed elements in the page <head> section:')
doc.add_paragraph('One for JSON-LD Code field (structured data)', style='List Bullet')
doc.add_paragraph('One for Hreflang Tags field (language alternates)', style='List Bullet')

# ── Step 3 ──
doc.add_heading('Step 3: Import Content via CSV (5 min)', level=2)
doc.add_paragraph('1. Go to CMS \u2192 Crypto Card Pages \u2192 click Import')
doc.add_paragraph('2. Upload cms_import.csv from the pages folder')
doc.add_paragraph('3. Map CSV columns to CMS fields:')

add_table(doc, ['CSV Column', 'CMS Field'], [
    ['Name', 'Name'],
    ['Slug', 'Slug'],
    ['SEO Title', 'SEO Title'],
    ['SEO Description', 'SEO Description'],
    ['H1', 'H1'],
    ['Hero Subtitle', 'Hero Subtitle'],
    ['Body Content', 'Body Content'],
    ['FAQ Content', 'FAQ Content'],
    ['Language', 'Language'],
    ['Country', 'Country'],
    ['Primary Keyword', 'Primary Keyword'],
    ['JSON-LD Code', 'JSON-LD Code'],
    ['Hreflang Tags', 'Hreflang Tags'],
])

doc.add_paragraph('4. Click Import \u2014 all ~41 pages created at once')

# ── Step 4 ──
doc.add_heading('Step 4: Review & Publish (5 min)', level=2)
doc.add_paragraph('1. Open a few pages in the Editor to verify content looks right:')
doc.add_paragraph('kolo.xyz/crypto-card/uae', style='List Bullet')
doc.add_paragraph('kolo.xyz/crypto-card/uk', style='List Bullet')
doc.add_paragraph('kolo.xyz/crypto-card/kripto-karta-oae (Russian)', style='List Bullet')
doc.add_paragraph('2. Check mobile responsiveness')
doc.add_paragraph('3. Publish all changes')

# ── Step 5 ──
doc.add_heading('Step 5: Post-Publish Checklist', level=2)
for item in [
    'Verify pages load: kolo.xyz/crypto-card/uae',
    'Check mobile layout',
    'Verify sitemap includes new pages: kolo.xyz/sitemap.xml',
    'Submit updated sitemap to Google Search Console (sc-domain:kolo.xyz)',
    'Spot-check JSON-LD with Google Rich Results Test',
]:
    doc.add_paragraph(item, style='List Bullet')

# ── Design Reference ──
doc.add_heading('Design Reference', level=2)
doc.add_paragraph('Match the existing kolo.xyz brand:')

add_table(doc, ['Element', 'Value'], [
    ['Background', '#0D0D0F (near black)'],
    ['Text', '#FFFFFF (white)'],
    ['Accent / CTA', '#6C5CE7 (purple)'],
    ['Font', 'Same as kolo.xyz (Inter or system)'],
    ['Max content width', '720px'],
    ['Button style', 'Rounded, purple background, white text'],
])

# ── Page Inventory ──
doc.add_heading('Page Inventory (~41 pages)', level=2)

doc.add_heading('English (~20 pages)', level=3)
add_table(doc, ['Slug', 'Target Keyword', 'Country'], [
    ['uae', 'crypto card UAE', 'UAE'],
    ['uk', 'crypto card UK', 'United Kingdom'],
    ['italy', 'crypto card Italy', 'Italy'],
    ['spain', 'crypto card Spain', 'Spain'],
    ['poland', 'crypto card Poland', 'Poland'],
    ['germany', 'crypto card Germany', 'Germany'],
    ['georgia', 'crypto card Georgia', 'Georgia'],
    ['cyprus', 'crypto card Cyprus', 'Cyprus'],
    ['latvia', 'crypto card Latvia', 'Latvia'],
    ['romania', 'crypto card Romania', 'Romania'],
    ['indonesia', 'crypto card Indonesia', 'Indonesia'],
    ['montenegro', 'crypto card Montenegro', 'Montenegro'],
    ['europe', 'crypto card Europe', 'Europe'],
    ['digital-nomads', 'crypto card digital nomads', 'Global'],
    ['freelancers', 'crypto card freelancers', 'Global'],
    ['business', 'crypto card for business', 'Global'],
    ['comparison-2026', 'crypto card comparison 2026', 'Global'],
    ['cashback', 'crypto card cashback bitcoin', 'Global'],
    ['usdt-card', 'USDT Visa card', 'Global'],
    ['low-fees', 'crypto card low fees', 'Global'],
])

doc.add_heading('Russian (~15 pages)', level=3)
add_table(doc, ['Slug', 'Target Keyword', 'Country'], [
    ['ru/oae', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u041e\u0410\u042d', 'UAE'],
    ['ru/ispaniya', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0418\u0441\u043f\u0430\u043d\u0438\u044f', 'Spain'],
    ['ru/gruziya', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0413\u0440\u0443\u0437\u0438\u044f', 'Georgia'],
    ['ru/kipr', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u041a\u0438\u043f\u0440', 'Cyprus'],
    ['ru/latviya', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u041b\u0430\u0442\u0432\u0438\u044f', 'Latvia'],
    ['ru/armeniya', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0410\u0440\u043c\u0435\u043d\u0438\u044f', 'Armenia'],
    ['ru/uzbekistan', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0423\u0437\u0431\u0435\u043a\u0438\u0441\u0442\u0430\u043d', 'Uzbekistan'],
    ['ru/kyrgyzstan', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u041a\u044b\u0440\u0433\u044b\u0437\u0441\u0442\u0430\u043d', 'Kyrgyzstan'],
    ['ru/azerbaydzhan', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0410\u0437\u0435\u0440\u0431\u0430\u0439\u0434\u0436\u0430\u043d', 'Azerbaijan'],
    ['ru/evropa', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0415\u0432\u0440\u043e\u043f\u0430', 'Europe'],
    ['ru/frilanserov', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0444\u0440\u0438\u043b\u0430\u043d\u0441\u0435\u0440\u043e\u0432', 'Global'],
    ['ru/biznes', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u0434\u043b\u044f \u0431\u0438\u0437\u043d\u0435\u0441\u0430', 'Global'],
    ['ru/usdt', 'USDT \u043a\u0430\u0440\u0442\u0430 Visa', 'Global'],
    ['ru/sravnenie', '\u0441\u0440\u0430\u0432\u043d\u0435\u043d\u0438\u0435 \u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442 2026', 'Global'],
    ['ru/keshbek', '\u043a\u0440\u0438\u043f\u0442\u043e \u043a\u0430\u0440\u0442\u0430 \u043a\u044d\u0448\u0431\u044d\u043a', 'Global'],
])

doc.add_heading('Italian (~3 pages)', level=3)
add_table(doc, ['Slug', 'Target Keyword'], [
    ['it/italia', 'carta crypto Italia'],
    ['it/visa-italia', 'carta crypto Visa Italia'],
    ['it/commissioni', 'carta crypto commissioni basse'],
])

doc.add_heading('Spanish (~3 pages)', level=3)
add_table(doc, ['Slug', 'Target Keyword'], [
    ['es/espana', 'tarjeta crypto Espa\u00f1a'],
    ['es/visa-espana', 'tarjeta crypto Visa Espa\u00f1a'],
    ['es/comisiones', 'tarjeta crypto comisiones bajas'],
])

# ── Notes ──
doc.add_heading('Important Notes', level=2)
p = doc.add_paragraph()
p.add_run('Do NOT create pages for: ').bold = True
p.add_run('US, Russia, Belarus, Turkey, Israel, Kazakhstan, China, India \u2014 Kolo cannot issue cards there')

doc.add_paragraph('kolo.in and kolo.xyz are the same site \u2014 kolo.in redirects to kolo.xyz', style='List Bullet')
doc.add_paragraph('All content is AI-generated and SEO-optimized \u2014 no editing needed unless you spot errors', style='List Bullet')
doc.add_paragraph('The CSV is generated by the SEO agent app at kolo-seo.streamlit.app \u2192 Programmatic SEO \u2192 Tab 4', style='List Bullet')
doc.add_paragraph('Questions? Contact the SEO team', style='List Bullet')

p = doc.add_paragraph()
p.add_run('\nGenerated by Kolo SEO Agent').italic = True

doc.save('/Users/ek/SEO agent/WEBFLOW_PROGRAMMATIC_SEO_BRIEF.docx')
print('Done')
